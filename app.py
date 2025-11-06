import os
import re
from collections import Counter

import PyPDF2
import torch
from docx import Document
from flask import Flask, redirect, render_template, request, url_for
from sentence_transformers import util
from werkzeug.utils import secure_filename

from model_loader import load_hf_model, load_local_models

COMMON_TECH_KEYWORDS = [
    "python",
    "java",
    "c++",
    "c#",
    "javascript",
    "typescript",
    "go",
    "rust",
    "ruby",
    "php",
    "swift",
    "kotlin",
    "tensorflow",
    "pytorch",
    "keras",
    "react",
    "angular",
    "vue",
    "django",
    "flask",
    "spring",
    "node",
    "express",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "sql",
    "postgres",
    "mysql",
    "mongodb",
    "redis",
    "spark",
    "hadoop",
    "nlp",
    "computer vision",
    "microservices",
    "rest",
    "graphql",
    "ci/cd",
    "terraform",
    "ansible",
    "linux",
]


def read_cv_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()

    try:
        if ext == ".pdf":
            with open(filepath, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = []
                for page in reader.pages:

                    page_text = page.extract_text() or ""
                    # Clean up common PDF extraction issues
                    page_text = re.sub(r"\s+", " ", page_text)
                    page_text = re.sub(r"[^\x20-\x7E\n]", "", page_text)
                    text.append(page_text.strip())

                extracted_text = "\n".join(text).strip()
                if not extracted_text:
                    raise ValueError("No text could be extracted from the PDF")
                return extracted_text

        elif ext == ".docx":
            doc = Document(filepath)
            # Extract both paragraphs and tables
            text_parts = []

            # Get text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # Get text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            return "\n".join(text_parts)

        elif ext == ".txt":
            with open(filepath, "r", encoding="utf-8") as file:
                text = file.read()
                if not text.strip():
                    raise ValueError("The text file is empty")
                return text

        else:
            raise ValueError(f"Unsupported file format: {ext}")

    except Exception as e:
        raise Exception(f"Failed to read file: {str(e)}")


def extract_technical_terms(text, qa_model, embedder, top_n=5):

    def clean_text(text):

        text = text.strip()
        text = re.sub(r"\s+", " ", text)
        return text

    try:

        questions = [
            "List all programming languages mentioned in the text.",
            "What specific technical skills or technologies are mentioned in the text?",
            "What software applications or development tools are listed in the text?",
            "What frameworks, libraries, or packages are mentioned in the text?",
            "What databases, storage systems, or data technologies are discussed in the text?",
        ]

        chunk_size = 512
        chunks = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]

        all_skills = set()

        for chunk in chunks:
            clean_chunk = clean_text(chunk)

            for question in questions:
                try:

                    result = qa_model(
                        question=question,
                        context=clean_chunk,
                        handle_impossible_answer=True,
                        max_answer_len=100,
                    )

                    if result and result.get("answer") and result.get("score", 0) > 0.1:
                        # Split the answer into individual terms
                        extracted_terms = re.split(
                            r"[,;\n]|\band\b|\bor\b|\bas well as\b|\bsuch as\b|\bincluding\b",
                            result["answer"].lower(),
                        )

                        # Clean and add terms
                        for term in extracted_terms:
                            term = term.strip("., ()[]{}")
                            if term and len(term) > 2:  # Avoid very short terms
                                # Remove common filler words
                                term = re.sub(
                                    r"\b(the|a|an|in|on|at|with|using|for|to|of)\b",
                                    "",
                                    term,
                                )
                                term = re.sub(r"\s+", " ", term).strip()
                                if term and len(term) > 2:
                                    all_skills.add(term)

                except Exception as e:
                    print(f"Extraction failed for chunk with question {question}: {e}")
                    continue

        text_lower = text.lower()
        for tech in COMMON_TECH_KEYWORDS:
            if re.search(rf"\b{re.escape(tech)}\b", text_lower):
                all_skills.add(tech.lower())

        unique_skills = list(dict.fromkeys(all_skills))

        if embedder and unique_skills:
            cv_embedding = embedder.encode(text_lower, convert_to_tensor=True)
            verified_skills = []

            for skill in unique_skills:

                skill_pattern = re.escape(skill)
                context_matches = re.finditer(skill_pattern, text_lower)
                contexts = []

                for match in context_matches:
                    start = max(0, match.start() - 100)
                    end = min(len(text), match.end() + 100)
                    contexts.append(text[start:end])

                if contexts:

                    context_embeddings = embedder.encode(
                        contexts, convert_to_tensor=True
                    )
                    similarities = util.cos_sim(cv_embedding, context_embeddings)
                    max_sim = float(similarities.max())

                    if max_sim > 0.4:
                        verified_skills.append((skill, max_sim))
                else:

                    skill_embedding = embedder.encode(skill, convert_to_tensor=True)
                    similarity = float(util.cos_sim(cv_embedding, skill_embedding))
                    if similarity > 0.4:
                        verified_skills.append((skill, similarity))

            verified_skills.sort(key=lambda x: x[1], reverse=True)
            return [skill for skill, _ in verified_skills[:top_n]]

        return unique_skills[:top_n]

    except Exception as e:
        print(f"Skill extraction failed: {e}")
        # Fallback to basic keyword matching
        text_lower = text.lower()
        matches = [
            tech
            for tech in COMMON_TECH_KEYWORDS
            if re.search(rf"\b{re.escape(tech)}\b", text_lower)
        ]
        return matches[:top_n]

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", os.urandom(24))
app.config["UPLOAD_FOLDER"] = "./uploads"
UPLOAD_FOLDER = os.path.join(os.getcwd(), "uploads")
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)


print("Loading models... (this may take a minute)")
flan_tokenizer, flan_model, embedder, keyword_extractor = load_local_models()
print("Models loaded successfully!")


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        if "cv_file" not in request.files:
            return render_template("index.html", error="No file uploaded.")

        file = request.files["cv_file"]
        if file.filename == "":
            return render_template("index.html", error="No file selected.")

        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
        file.save(filepath)

        try:
            cv_text = read_cv_file(filepath)
            if not cv_text.strip():
                return render_template(
                    "index.html", error="Could not extract text from the uploaded file."
                )

            techs = extract_technical_terms(
                cv_text, keyword_extractor, embedder, top_n=5
            )

            if not techs:
                return render_template(
                    "index.html",
                    error="Could not identify technical skills in the CV. Please ensure your CV includes technical terms.",
                )

            tech_summary = ", ".join(techs)
        except Exception as e:
            return render_template("index.html", error=f"Failed to process CV: {e}")

        return render_template(
            "interview.html", cv_text=cv_text, tech_summary=tech_summary
        )
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate_questions():
    cv_text = request.form["cv_text"]
    tech_summary = request.form["tech_summary"]
    question_type = request.form.get(
        "question_type", "technical"
    )  # Default to technical if not specified

    persona = "You are a sharp, experienced technical interviewer conducting an important job interview."

    role = persona + "\n"

    if question_type == "technical":
        role += f"""Focus area: {tech_summary}
Your task: Generate a specific technical question that:
1. Relates to the candidate's actual experience from their CV
2. Probes their deep understanding of {tech_summary}
3. Asks about a real technical challenge they faced"""
    elif question_type == "soft_skills":
        role += """Focus area: Professional Behavior and Team Dynamics
Your task: Generate a behavioral question that:
1. Explores how they handle workplace challenges
2. Reveals their approach to teamwork and communication
3. Provides insight into their problem-solving style"""
    else:  # motivation
        role += f"""Focus area: Career Goals and Motivation
Your task: Generate a question that:
1. Explores their career trajectory and aspirations
2. Reveals their motivation for working with {tech_summary}
3. Uncovers what drives them professionally"""

    base_context = f"""Analyze this CV excerpt carefully:
{cv_text[:2000]}

Key technical skills: {tech_summary}
"""

    if question_type == "technical":
        prompt = f"""{base_context}

You are an expert technical interviewer. Generate a challenging but fair technical interview question that:
1. Focuses specifically on their experience with {tech_summary}
2. Tests both theoretical understanding and practical application
3. Encourages them to describe a specific technical challenge they solved
4. Reveals their problem-solving approach and technical depth

The question should:
- Be specific to their actual experience, not generic
- Focus on complex problem-solving, not just tool usage
- Encourage detailed technical discussion
- Be answerable with a concrete example from their work

Format: Generate only the question, starting with "Can you describe" or "Tell me about"
Question:"""

    elif question_type == "soft_skills":
        prompt = f"""{base_context}

You are an experienced behavioral interviewer. Generate a situational interview question that:
1. Explores how they've applied {tech_summary} in team settings
2. Reveals their approach to technical collaboration
3. Uncovers their communication style with both technical and non-technical stakeholders
4. Shows how they handle challenges and conflicts

The question should:
- Focus on real scenarios from their experience
- Encourage STAR method responses
- Reveal both technical and interpersonal skills
- Target specific behavioral competencies

Format: Generate only the question, starting with "Tell me about" or "Describe a time"
Question:"""

    else:  # motivation
        prompt = f"""{base_context}

You are an insightful career development interviewer. Generate a motivation-focused question that:
1. Explores their journey with {tech_summary}
2. Reveals their technical growth mindset
3. Uncovers their long-term career vision
4. Shows what drives their technical choices

The question should:
- Connect their past experiences to future goals
- Reveal their passion for technology
- Explore their technical decision-making
- Uncover their professional values

Format: Generate only the question, starting with "What motivated" or "How has"
Question:"""

    inputs = flan_tokenizer(
        prompt, return_tensors="pt", truncation=True, max_length=1024
    ).to(flan_model.device)
    outputs = flan_model.generate(
        **inputs,
        max_new_tokens=100,
        do_sample=True,
        temperature=0.7,
        top_p=0.9,
        num_return_sequences=1,
    )
    question = flan_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

    # Clean up the generated question
    question = re.sub(r"[\[\]\{\}]", "", question)
    
    if "Question:" in question:
        question = question.split("Question:")[-1].strip()

    question = re.sub(
        r"^(Here'?s? (?:a |an )?(?:possible |suggested |example )?question:?\s*)",
        "",
        question,
        flags=re.IGNORECASE,
    )

    if not question:
        question = "Could you tell me about your experience?"
    else:
        question = (
            question.rstrip(".") + "?" if not question.endswith("?") else question
        )
        question = question[0].upper() + question[1:]

    return render_template(
        "interview.html",
        cv_text=cv_text,
        tech_summary=tech_summary,
        questions=[question],
    )

@app.route("/evaluate", methods=["POST"])
def evaluate_answer():
    question = request.form["question"]
    user_answer = request.form["user_answer"]
    tech_summary = request.form["tech_summary"]

    if len(user_answer.strip()) < 10:
        return render_template(
            "interview.html",
            question=question,
            user_answer=user_answer,
            feedback="Please provide a more detailed answer to receive meaningful feedback.",
            quality_score=0,
            tech_summary=tech_summary,
        )

    ideal_prompt = f"""Question: {question}

Create a model STAR method answer that:
1. Describes a specific Situation
2. Explains the Task clearly
3. Details the Actions taken
4. Shows measurable Results
5. Demonstrates expertise in {tech_summary}

Answer in 2-3 paragraphs."""

    inputs = flan_tokenizer(
        ideal_prompt, return_tensors="pt", truncation=True, max_length=1024
    ).to(flan_model.device)
    outputs = flan_model.generate(
        **inputs, max_new_tokens=200, num_beams=2, temperature=0.3
    )
    ideal_answer = flan_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()

    try:
        embeddings = embedder.encode(
            [user_answer, ideal_answer], convert_to_tensor=True
        )
        similarity_score = float(util.cos_sim(embeddings[0], embeddings[1]).item())
    except Exception:
        similarity_score = 0.0

    star_prompt = f"""Rate this answer on STAR method implementation (0-10):
Question: {question}
Answer: {user_answer}
Rate only the presence and quality of: Situation, Task, Action, Results
Score:"""

    tech_prompt = f"""Rate this answer on technical depth (0-10):
Expected expertise in: {tech_summary}
Answer: {user_answer}
Rate only technical accuracy and depth.
Score:"""

    # Get STAR and Technical scores
    star_score = generate_score(flan_model, flan_tokenizer, star_prompt)
    tech_score = generate_score(flan_model, flan_tokenizer, tech_prompt)

    # Calculate final scores
    star_score = min(10, max(0, star_score))
    tech_score = min(10, max(0, tech_score))
    quality_score = (star_score + tech_score + similarity_score * 10) / 30.0

    # Generate focused feedback based on scores
    feedback_sections = []

    # STAR Method Feedback
    feedback_sections.append(f"STAR Method: {star_score}/10")
    if star_score < 5:
        feedback_sections.append(
            "-Missing key STAR components - structure your answer with Situation, Task, Action, and Results"
        )
    elif star_score < 8:
        feedback_sections.append(
            "-Good attempt at STAR format but needs more specific details"
        )
    else:
        feedback_sections.append("-Excellent use of STAR method")

    # Technical Feedback
    feedback_sections.append(f"\nTechnical Skills: {tech_score}/10")
    if tech_score < 5:
        feedback_sections.append(
            f"-Include more specific examples of using {tech_summary}"
        )
    elif tech_score < 8:
        feedback_sections.append("-Good technical content, consider adding more depth")
    else:
        feedback_sections.append("-Strong technical demonstration")

    # Overall Quality
    similarity_percent = int(similarity_score * 100)
    feedback_sections.append(f"\nResponse Completeness: {similarity_percent}%")
    if similarity_percent < 50:
        feedback_sections.append("-Answer needs more detail and examples")
    elif similarity_percent < 80:
        feedback_sections.append("-Good response, could use more specific outcomes")
    else:
        feedback_sections.append("-Very comprehensive answer")

    # Specific Improvements
    feedback_sections.append("\nSuggested Improvements:")
    missing_elements = []
    if "situation" not in user_answer.lower():
        missing_elements.append(
            "-Set the context by describing the specific situation"
        )
    if "result" not in user_answer.lower() and "outcome" not in user_answer.lower():
        missing_elements.append("-Include quantifiable results or outcomes")
    if not any(
        tech.lower() in user_answer.lower() for tech in tech_summary.split(", ")
    ):
        missing_elements.append(f"-Mention specific examples of using {tech_summary}")
    if not missing_elements:
        missing_elements.append(
            "-Add more quantifiable metrics to strengthen your response"
        )
    feedback_sections.extend(missing_elements)

    feedback = "\n".join(feedback_sections)

    return render_template(
        "interview.html",
        question=question,
        user_answer=user_answer,
        feedback=feedback,
        quality_score=round(quality_score, 2),
        tech_summary=tech_summary,
    )



