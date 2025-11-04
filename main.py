import os
import random
import re
import sys
from collections import Counter

import torch
from sentence_transformers import SentenceTransformer, util
from transformers import (
    AutoModelForSeq2SeqLM,
    AutoModelForTokenClassification,
    AutoTokenizer,
    pipeline,
)

from model_loader import load_hf_model, load_local_models

COMMON_TECH_KEYWORDS = [
    "python",
    "java",
    "c++",
    "c#",
    "javascript",
    "typescript",
    "go",
    "rust",
    "ruby",
    "php",
    "swift",
    "kotlin",
    "tensorflow",
    "pytorch",
    "keras",
    "react",
    "angular",
    "vue",
    "django",
    "flask",
    "spring",
    "node",
    "express",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "gcp",
    "sql",
    "postgres",
    "mysql",
    "mongodb",
    "redis",
    "spark",
    "hadoop",
    "nlp",
    "computer vision",
    "microservices",
    "rest",
    "graphql",
    "ci/cd",
    "terraform",
    "ansible",
    "linux",
]


def extract_technical_terms(text, keyword_extractor, top_n=5):

    try:
        ner_results = keyword_extractor(text)
        raw_keywords = [
            entity["word"].lower().strip()
            for entity in ner_results
            if entity["score"] > 0.3
        ]

        found_tech = []
        for kw in raw_keywords:
            for tech in COMMON_TECH_KEYWORDS:
                if tech in kw or kw in tech:
                    found_tech.append(tech)
                    break

        unique_tech = list(dict.fromkeys(found_tech))
        return unique_tech[:top_n]

    except Exception as e:
        print(f"NER extraction failed (falling back to regex): {e}")
        text_lower = text.lower()
        found = [
            kw
            for kw in COMMON_TECH_KEYWORDS
            if re.search(rf"\b{re.escape(kw)}\b", text_lower)
        ]
        freq = Counter(re.findall(r"\w[\w\+\#\-/\.]*", text_lower))
        return sorted(found, key=lambda k: freq.get(k, 0), reverse=True)[:top_n]


def main():
    try:
        flan_tokenizer, flan_model, embedder, keyword_extractor = load_local_models()
    except ValueError as e:
        print(f"FATAL ERROR: Expected 4 models, got mismatch: {e}")
        sys.exit(1)

    # Create uploads directory if it doesn't exist
    uploads_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
    os.makedirs(uploads_dir, exist_ok=True)

    # Check for CV file in uploads directory
    cv_files = [
        f for f in os.listdir(uploads_dir) if f.endswith((".txt", ".pdf", ".docx"))
    ]

    if not cv_files:
        print("ERROR: No CV files found in /uploads directory.")
        print("Please place your CV file (PDF, DOCX, or TXT) in the 'uploads' folder.")
        return

    if len(cv_files) > 1:
        print("\nMultiple CV files found in uploads directory:")
        for idx, file in enumerate(cv_files, 1):
            print(f"  {idx}. {file}")

        while True:
            try:
                choice = input("\nEnter the number of the CV to use (or 'q' to quit): ")
                if choice.lower() == "q":
                    return

                file_idx = int(choice) - 1
                if 0 <= file_idx < len(cv_files):
                    cv_file = cv_files[file_idx]
                    break
                else:
                    print("Invalid selection. Please try again.")
            except ValueError:
                print("Please enter a valid number.")
    else:
        cv_file = cv_files[0]

    # Get the full path to the CV file
    cv_file = os.path.join(uploads_dir, cv_file)
    print(f"\nUsing CV: {os.path.basename(cv_file)}")
    ext = os.path.splitext(cv_file)[1].lower()
    print("\nReading CV file...")

    try:
        if ext == ".pdf":
            try:
                import PyPDF2

                with open(cv_file, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    text = []
                    for page in reader.pages:
                        page_text = page.extract_text() or ""
                        # Clean up common PDF extraction issues
                        page_text = re.sub(
                            r"\s+", " ", page_text
                        )  # normalize whitespace
                        page_text = re.sub(
                            r"[^\x20-\x7E\n]", "", page_text
                        )  # remove non-printable chars
                        text.append(page_text.strip())
                    cv_text = "\n".join(text).strip()
                print("✓ Successfully read PDF file")
            except ImportError:
                print("ERROR: PyPDF2 not installed. Run: pip install PyPDF2")
                return

        elif ext == ".docx":
            try:
                from docx import Document

                doc = Document(cv_file)
                # Extract both paragraphs and tables
                text_parts = []

                # Get text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())

                # Get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " ".join(
                            cell.text.strip() for cell in row.cells if cell.text.strip()
                        )
                        if row_text:
                            text_parts.append(row_text)

                cv_text = "\n".join(text_parts)
                print("✓ Successfully read DOCX file")
            except ImportError:
                print("ERROR: python-docx not installed. Run: pip install python-docx")
                return

        elif ext == ".txt":
            with open(cv_file, "r", encoding="utf-8") as file:
                cv_text = file.read()
            print("✓ Successfully read TXT file")
        else:
            print(f"ERROR: Unsupported file format: {ext}")
            print("Please provide a PDF, DOCX, or TXT file.")
            return

        if not cv_text.strip():
            print(f"ERROR: No text could be extracted from {os.path.basename(cv_file)}")
            print("Please make sure the file contains extractable text.")
            return

    except Exception as e:
        print(f"ERROR: Failed to read {cv_file}: {str(e)}")
        return

    detected_techs = extract_technical_terms(cv_text, keyword_extractor, top_n=5)
    tech_summary = ", ".join(detected_techs) if detected_techs else "programming and IT"
    print(f"Detected technical keywords: {tech_summary}\n")

    asked_questions = set()
    all_answers = []

    # Generate 2 questions of each type
    question_types = [
        "technical",
        "technical",
        "soft_skills",
        "soft_skills",
        "motivation",
        "motivation",
    ]

    for i, q_type in enumerate(question_types, 1):
        persona = "You are a sharp, creative, and insightful interviewer."

        if q_type == "technical":
            angle = random.choice(
                [
                    "ask about a complex debugging scenario they faced.",
                    "ask about a key design decision and its trade-offs.",
                    "ask about optimizing performance in a past project.",
                    "ask about scalability challenges they've encountered.",
                    f"ask about a time they had to learn a new tool from {tech_summary} quickly.",
                    "ask about their experience with code reviews and testing.",
                ]
            )
            role = (
                f"{persona} You are a senior technical interviewer focusing on {tech_summary}. "
                f"Your task: {angle}"
            )
        elif q_type == "soft_skills":
            angle = random.choice(
                [
                    "ask a behavioral question about handling difficult feedback.",
                    "ask a situational question about managing conflicting priorities.",
                    "ask about a time they had to persuade a team member to their point of view.",
                    "ask about a project failure and what they learned from it.",
                    "ask about a time they resolved a team conflict.",
                ]
            )
            role = f"{persona} You are an HR interviewer. Your task: {angle}"
        else:
            angle = random.choice(
                [
                    "ask about their long-term career vision (5+ years).",
                    "ask what kind of work environment helps them thrive.",
                    "ask about a project from their CV that they found most motivating and why.",
                    "ask how they stay up-to-date with industry trends.",
                    "ask about a time they went above and beyond their assigned role.",
                ]
            )
            role = f"{persona} You are a recruiter. Your task: {angle}"

        prompt = f"""
CV Context:
{cv_text}

Interviewer Role and Task: {role}

Rules for your generated question:
1.  Generate ONE SINGLE interview question.
2.  The question MUST be based on the Role, Task, and CV Context.
3.  DO NOT use generic, common questions (e.g., "What is your greatest weakness?", "Tell me about yourself.").
4.  The question must be unique, specific, and insightful.
5.  Start with a capital letter and end with a single question mark '?'.
6.  Do not add any preamble like "Here is your question:".

Question:
"""

        print(f"\n--- Generating {q_type.capitalize()} Question (Angle: {angle}) ---")
        question = None
        for _ in range(3):
            try:
                inputs = flan_tokenizer(
                    prompt, return_tensors="pt", truncation=True, max_length=1024
                ).to(flan_model.device)

                outputs = flan_model.generate(
                    **inputs,
                    max_new_tokens=80,
                    do_sample=True,
                    temperature=0.85,
                    top_k=50,
                    top_p=0.95,
                    num_beams=1,
                    no_repeat_ngram_size=2,
                )

                q = flan_tokenizer.decode(outputs[0], skip_special_tokens=True).strip()
                q = re.sub(r"[\[\]\{\}]", "", q)
                if q and q.endswith("?") and len(q) > 15 and q not in asked_questions:
                    question = q
                    break
            except Exception as e:
                print(f"Generation failed: {e}")

        if not question:
            question = f"Tell me about a challenging {q_type} situation you faced."
            print("Used fallback question.")

        asked_questions.add(question)
        print(f"\nQuestion {i}/6: {question}")

        user_answer = input("\nYour answer (in English): ").strip()

        # Generate ideal English answer
        ideal_prompt = f"""
Question: {question}

Provide a strong, concise English answer using the STAR method (Situation, Task, Action, Result).
Use specific tools ({tech_summary}), metrics, and outcomes.
Be concise (1-2 paragraphs).
"""
        try:
            inputs = flan_tokenizer(
                ideal_prompt, return_tensors="pt", truncation=True, max_length=1024
            ).to(flan_model.device)
            # Use simple generation for the ideal answer, no need for complex sampling
            outputs = flan_model.generate(**inputs, max_new_tokens=200, num_beams=2)
            ideal_english = flan_tokenizer.decode(
                outputs[0], skip_special_tokens=True
            ).strip()

        except Exception as e:
            ideal_english = "(Ideal answer generation failed)"

        # Score similarity
        score = 0.0
        if user_answer and ideal_english and "failed" not in ideal_english:
            try:
                embeddings = embedder.encode(
                    [user_answer, ideal_english], convert_to_tensor=True
                )
                score = util.cos_sim(embeddings[0], embeddings[1]).item()
            except Exception as e:
                print(f"Embedding/scoring failed: {e}")
                score = 0.0

        all_answers.append(
            {
                "question": question,
                "user_answer": user_answer,
                "ideal_answer": ideal_english,
                "score": score,
            }
        )

        # Feedback
        print("\n=== Feedback ===")
        print(f"Ideal: {ideal_english}")
        print(f"Score: {score:.2%}")
        if score > 0.8:
            print("Outstanding!")
        elif score > 0.6:
            print("Solid, but add more detail.")
        else:
            print("Needs more structure and specifics.")

        if i < 6:
            input("\nPress Enter for next question...")

    # Final Summary
    print("\n" + "=" * 60)
    print("FINAL INTERVIEW REPORT")
    print("=" * 60)

    # Calculate scores by category
    technical_scores = [a["score"] for i, a in enumerate(all_answers) if i < 2]
    soft_scores = [a["score"] for i, a in enumerate(all_answers) if 2 <= i < 4]
    motivation_scores = [a["score"] for i, a in enumerate(all_answers) if i >= 4]

    avg_technical = (
        sum(technical_scores) / len(technical_scores) if technical_scores else 0
    )
    avg_soft = sum(soft_scores) / len(soft_scores) if soft_scores else 0
    avg_motivation = (
        sum(motivation_scores) / len(motivation_scores) if motivation_scores else 0
    )
    overall_avg = (avg_technical + avg_soft + avg_motivation) / 3

    # Print questions by category
    print("\nTechnical Questions:")
    for i in range(2):
        print(
            f"  Q{i+1} [{all_answers[i]['score']:.1%}]: {all_answers[i]['question'][:60]}..."
        )
    print(f"  Average: {avg_technical:.1%}")

    print("\nBehavioral Questions:")
    for i in range(2, 4):
        print(
            f"  Q{i+1} [{all_answers[i]['score']:.1%}]: {all_answers[i]['question'][:60]}..."
        )
    print(f"  Average: {avg_soft:.1%}")

    print("\nMotivation Questions:")
    for i in range(4, 6):
        print(
            f"  Q{i+1} [{all_answers[i]['score']:.1%}]: {all_answers[i]['question'][:60]}..."
        )
    print(f"  Average: {avg_motivation:.1%}")

    print("\n" + "-" * 60)
    print(f"Overall Score: {overall_avg:.1%}")

    # Overall assessment
    strengths = []
    areas_to_improve = []

    if avg_technical > 0.7:
        strengths.append("Technical knowledge")
    elif avg_technical < 0.5:
        areas_to_improve.append("Technical depth")

    if avg_soft > 0.7:
        strengths.append("Behavioral competencies")
    elif avg_soft < 0.5:
        areas_to_improve.append("STAR method responses")

    if avg_motivation > 0.7:
        strengths.append("Career clarity")
    elif avg_motivation < 0.5:
        areas_to_improve.append("Professional motivation")

    if strengths:
        print("\nStrengths:", ", ".join(strengths))
    if areas_to_improve:
        print("Areas to Improve:", ", ".join(areas_to_improve))

    print("\nOverall Assessment:")
    if overall_avg > 0.75:
        print("Strong candidate! Ready for advanced interviews.")
    elif overall_avg > 0.6:
        print("Good potential with some preparation needed.")
    elif overall_avg > 0.4:
        print("More practice recommended, focus on identified areas.")
    else:
        print("Significant preparation needed before formal interviews.")


if __name__ == "__main__":
    main()
